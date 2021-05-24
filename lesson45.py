# reading and editing a word document

# Document object contains Paragraph objects
# Paragraphs objects contain smaller Run objects

import docx

# d = docx.Document('/Users/jordannewberry/Documents/demo.docx')

# print(d.paragraphs[0].text)
# print(d.paragraphs[1].text)
# print(d.paragraphs[2].text)
# print(d.paragraphs[3].text)
# print(d.paragraphs[4].text)
# print(d.paragraphs[5].text)



# p = d.paragraphs[2]

# print(p.runs[0].text)
# print(p.runs[1].bold)
# print(p.runs[2].text)
# print(p.runs[3].italic)

# re-assign run variables
# p.runs[3].underline = True
# p.runs[3].text = 'italic and underlined.'
# d.save('/Users/jordannewberry/Documents/demo2.docx')



# print(p.runs[1].text)
# print(p.runs[1].text)
# print(p.runs[2].text)


# p.style = 'Title' # this doesn't work in MacOS
# d.save('demo3.docx')


# creating new docx file, this only exists in the memory not the hard drive yet
d = docx.Document()

d.add_paragraph('Hello this is a paragraph.')

d.add_paragraph('This is another paragraph.')

# d.save('/Users/jordannewberry/Documents/demo4.docx')

p = d.paragraphs[0]

p.add_run(' This is a new run.')

p.runs[1].bold = True

d.save('/Users/jordannewberry/Documents/demo5.docx')

